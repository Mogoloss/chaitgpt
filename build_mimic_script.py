# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor


OUT = Path(r"C:\Users\mgq56\Desktop\生物进化的历程\生物进化的历程_逐字稿_仿写版.docx")


def set_run(run, font="宋体", size=12, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_normal(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run(r)
    return p


def add_colored_label(doc, label, text="", color=RGBColor(0, 112, 192)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run(label)
    set_run(r1, bold=True, color=color)
    if text:
        r2 = p.add_run(text)
        set_run(r2)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = title.add_run("《生物进化的历程》课堂教学逐字稿")
set_run(r, font="黑体", size=18, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = sub.add_run("人教版  八年级生物下册  第六单元·第三章·第二节")
set_run(r, size=11)

legend = doc.add_paragraph()
legend.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
r = legend.add_run("颜色说明：  ")
set_run(r, size=11)
r = legend.add_run("■ 蓝色【课堂提问】")
set_run(r, size=11, bold=True, color=RGBColor(0, 112, 192))
r = legend.add_run("    ")
set_run(r, size=11)
r = legend.add_run("■ 绿色【小组讨论】")
set_run(r, size=11, bold=True, color=RGBColor(0, 176, 80))
r = legend.add_run("    ")
set_run(r, size=11)
r = legend.add_run("■ 橙色【活动指导】")
set_run(r, size=11, bold=True, color=RGBColor(230, 145, 56))
r = legend.add_run("    ")
set_run(r, size=11)
r = legend.add_run("■ 紫色【预设回答】")
set_run(r, size=11, bold=True, color=RGBColor(112, 48, 160))

add_normal(doc, "一、课堂导入（约3分钟）", first_indent=False)
add_normal(doc, "【PPT第1—3张——导入与化石概念】", first_indent=False)
add_normal(doc, "同学们，上课！好，请坐。今天我们来学习第二节《生物进化的历程》。看到这个题目，大家先不要急着记结论，老师想先带大家从一个问题出发。")
add_normal(doc, "请大家看屏幕上的图片，这是一种生活在大约1.31亿年前的动物，叫郑氏始孔子鸟。它个头和鸡差不多，看起来有点像今天的鸟类，但它又保留了一些比较原始的特征。")
add_colored_label(doc, "【课堂提问】", "像这样一种早已灭绝的动物，科学家是怎样知道它曾经存在过的？")
add_colored_label(doc, "【预设回答】", "通过化石。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好，正是通过化石，科学家才能研究这些远古生物。那什么是化石呢？请大家看下一张。")
add_normal(doc, "郑氏始孔子鸟之所以在灭绝上亿年后还能为人所知，就是因为在地层中留下了化石。化石是指通过自然作用保存在地层中的古代生物的遗体、遗物或生活痕迹等。")
add_colored_label(doc, "【课堂提问】", "大家注意，化石只包括古代生物的遗体吗？")
add_colored_label(doc, "【预设回答】", "不是，还包括遗物和生活痕迹。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好。所以今天这节课，我们就围绕两个核心问题展开：第一，为什么说化石是研究生物进化的直接证据；第二，生物进化的大致历程究竟是怎样的。")

add_normal(doc, "二、探究新知（一）：研究生物进化的直接证据——化石（约12分钟）", first_indent=False)
add_normal(doc, "【PPT第4—14张——化石、地层与结论】", first_indent=False)
add_normal(doc, "下面我们先进入第一部分内容：研究生物进化的直接证据——化石。")
add_normal(doc, "请大家看课件中的几组化石材料。第一组是三叶虫化石。三叶虫是一类远古海洋生物，在很多地方都发现了它们的化石。")
add_colored_label(doc, "【课堂提问】", "三叶虫化石在很多地方都有发现，这说明了什么？")
add_colored_label(doc, "【预设回答】", "说明三叶虫曾经广泛生活在海洋中，也说明有些地方过去可能是海洋环境。", color=RGBColor(112, 48, 160))
add_normal(doc, "回答得很好。也就是说，化石不仅能告诉我们古代生物存在过，还能帮助我们推测当时的生活环境。")
add_normal(doc, "再看第二组材料，辽宁古果化石。这是一种原始被子植物化石，它把被子植物出现的时间进一步提前了。")
add_colored_label(doc, "【课堂提问】", "辽宁古果化石的发现，对研究植物进化有什么意义？")
add_colored_label(doc, "【预设回答】", "为研究被子植物的起源提供了新的化石证据。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好，科学研究正是在新证据不断出现的过程中向前推进的。")
add_normal(doc, "接着看胡氏耀龙化石。胡氏耀龙既有羽毛，又有牙和爪，看上去有点像鸟，但科学家经过分析比较后，认定它属于恐龙。")
add_colored_label(doc, "【课堂提问】", "胡氏耀龙为什么既像鸟又被认定属于恐龙？")
add_colored_label(doc, "【预设回答】", "因为它同时具有鸟类和恐龙的某些特征，属于过渡类型。", color=RGBColor(112, 48, 160))
add_normal(doc, "非常好。像胡氏耀龙这种兼具两类生物部分特征的类型，能够帮助我们理解不同类群之间的演变关系。")
add_normal(doc, "好，接下来请大家把注意力从单个化石，转向化石在地层中的分布。")
add_normal(doc, "一般来说，先沉积形成的地层在下面，后沉积形成的地层在上面。也就是说，下面的地层更古老，上面的地层更晚近。")
add_colored_label(doc, "【课堂提问】", "根据不同地层中的化石分布，你能发现什么规律？")
add_colored_label(doc, "【预设回答】", "越古老的地层中，化石越简单、越低等；越晚近的地层中，化石越复杂、越高等。", color=RGBColor(112, 48, 160))
add_normal(doc, "对，这个规律非常重要。它说明生物并不是一下子就变成今天这样的，而是在漫长的年代里不断变化、逐步演化而来的。")
add_normal(doc, "所以，到这里我们可以得出本节课的第一个核心结论：化石是研究生物进化最直接、最重要的证据。")

add_normal(doc, "三、科学方法：比较（约4分钟）", first_indent=False)
add_normal(doc, "【PPT第15张——科学方法：比较】", first_indent=False)
add_normal(doc, "有了证据之后，科学家怎样从这些证据中得出结论呢？这时候就要用到一种重要的科学方法——比较。")
add_normal(doc, "所谓比较，就是根据一定标准，把彼此有联系的事物放在一起，对照分析，从而找出它们之间的共同规律和本质联系。")
add_normal(doc, "比如课件中提到，比较马的前肢、鹰的翼和蝙蝠的前肢骨骼，我们会发现它们虽然功能不同，但骨骼排列有许多共同点。")
add_colored_label(doc, "【课堂提问】", "比较这种方法，在研究生物进化时有什么作用？")
add_colored_label(doc, "【预设回答】", "可以根据不同生物之间的相同点和不同点，推测它们的亲缘关系和进化关系。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好。也就是说，化石提供证据，比较帮助我们利用证据得出结论。前面两个问题我们已经解决了，下面进入第三个问题：生物进化的大致历程。")

add_normal(doc, "四、探究新知（二）：生物进化的大致历程（约14分钟）", first_indent=False)
add_normal(doc, "【PPT第16—21张——植物、动物历程与总结】", first_indent=False)
add_normal(doc, "前面我们已经明确了研究生物进化的证据是化石、方法是比较。下面我们就利用这些证据和方法，进一步看看生物进化的大致历程。")
add_normal(doc, "这部分课件的顺序是先看植物，再看无脊椎动物，再看脊椎动物，最后借助进化树和总结页，把整体规律提炼出来。")

add_normal(doc, "（一）植物进化的总体历程")
add_normal(doc, "最早的植物生活在原始海洋中，原始类型是藻类植物。随后，一部分植物逐渐向陆地过渡，先出现了苔藓植物，再出现了蕨类植物。")
add_colored_label(doc, "【课堂提问】", "苔藓植物和蕨类植物虽然已经生活在陆地上，但为什么说它们还没有真正适应陆地生活？")
add_colored_label(doc, "【预设回答】", "因为它们的生殖过程还离不开水。", color=RGBColor(112, 48, 160))
add_normal(doc, "对。再往后，出现了裸子植物，最后发展到被子植物。被子植物的生殖进一步摆脱了对水的依赖，因此成为今天分布最广的一类植物。")
add_normal(doc, "所以植物进化的大致历程可以概括为：原始藻类植物、原始苔藓植物、原始蕨类植物、原始裸子植物、原始被子植物。")
add_normal(doc, "从这条路线中大家可以清楚看到，植物进化的一个重要方向，就是不断增强对陆地环境的适应能力。植物讲完之后，我们再来看动物。")

add_normal(doc, "（二）无脊椎动物的进化历程")
add_normal(doc, "先看无脊椎动物。课件中给出的顺序是：原始单细胞动物、原始刺胞动物、原始扁形动物、原始线虫动物、原始软体动物、原始环节动物，最后发展到原始节肢动物。")
add_colored_label(doc, "【课堂提问】", "无脊椎动物的进化历程，最明显体现了哪一个趋势？")
add_colored_label(doc, "【预设回答】", "由简单到复杂。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好。随着进化，动物的身体结构越来越复杂，器官分化越来越明显。")
add_normal(doc, "也就是说，无脊椎动物这部分最能帮助我们理解‘结构由简单到复杂’这一趋势。接下来，再看脊椎动物。")

add_normal(doc, "（三）脊椎动物的进化历程")
add_normal(doc, "最后看脊椎动物。课件中展示得很清楚：脊椎动物大致经历了原始鱼类、原始两栖类、原始爬行类，然后由爬行类分别发展出原始鸟类和原始哺乳类。")
add_colored_label(doc, "【课堂提问】", "脊椎动物从鱼类到两栖类，再到爬行类、鸟类和哺乳类，这个过程体现了哪些进化趋势？")
add_colored_label(doc, "【预设回答】", "体现了由水生到陆生，也体现了由低等到高等、由简单到复杂。", color=RGBColor(112, 48, 160))
add_normal(doc, "非常好。课件最后还特别强调了一点：各种生物在进化过程中形成了各自适应环境的形态结构和生活习性。大家在理解进化时，一定要和‘适应环境’联系起来。")
add_normal(doc, "好，到这里植物、无脊椎动物和脊椎动物三条主要线索我们都梳理完了。下面请大家再看进化树和总结页，把这些零散知识连成整体。")

add_normal(doc, "（四）借助进化树和总结页把握整体规律")
add_normal(doc, "进化树帮助我们从整体上看到，不同生物类群并不是孤立存在的，而是在长期进化过程中不断分支、不断分化形成的。")
add_normal(doc, "而总结页则把本节课最核心的内容进行了提炼：研究生物进化的证据是化石，研究生物进化的方法是比较，生物进化具有三个总体趋势。")
add_colored_label(doc, "【课堂提问】", "谁能结合进化树和总结页，再完整说一遍生物进化的三个总体趋势？")
add_colored_label(doc, "【预设回答】", "结构方面由简单到复杂，生活环境方面由水生到陆生，进化水平方面由低等到高等。", color=RGBColor(112, 48, 160))

add_normal(doc, "五、课堂总结（约4分钟）", first_indent=False)
add_normal(doc, "【PPT第21张——总结】", first_indent=False)
add_normal(doc, "现在我们一起把整节课的核心内容再回顾一遍。")
add_colored_label(doc, "【课堂提问】", "研究生物进化的证据是什么？")
add_colored_label(doc, "【预设回答】", "化石。", color=RGBColor(112, 48, 160))
add_colored_label(doc, "【课堂提问】", "研究生物进化的方法是什么？")
add_colored_label(doc, "【预设回答】", "比较。", color=RGBColor(112, 48, 160))
add_colored_label(doc, "【课堂提问】", "生物进化的总体趋势是什么？")
add_colored_label(doc, "【预设回答】", "结构方面由简单到复杂，生活环境方面由水生到陆生，进化水平方面由低等到高等。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好，这三问其实就是本节课最核心的三句话，大家一定要记牢。")

add_normal(doc, "六、随堂练习与巩固（约5分钟）", first_indent=False)
add_normal(doc, "【PPT第22—23张——随堂练习】", first_indent=False)
add_normal(doc, "下面我们用两道练习题来检查一下学习效果。")
add_colored_label(doc, "【活动指导】", "请同学们先独立判断第23页和第24页的题目，再和同桌交换意见。")
add_colored_label(doc, "【课堂提问】", "为什么“化石是研究生物进化的唯一证据”这句话是错误的？")
add_colored_label(doc, "【预设回答】", "因为化石是最直接的证据，但不是唯一证据。", color=RGBColor(112, 48, 160))
add_colored_label(doc, "【课堂提问】", "为什么A项“从植物到动物”不属于生物进化的总体趋势？")
add_colored_label(doc, "【预设回答】", "因为植物和动物是不同的进化分支，不能简单理解为从植物进化到动物。", color=RGBColor(112, 48, 160))
add_normal(doc, "很好，通过这两道题，大家要特别注意三个易错点：第一，化石是最直接的证据，但不是唯一证据；第二，生物进化不能简单理解成‘从植物到动物’；第三，三个总体趋势一定要准确表述。")

add_normal(doc, "七、结束语与作业（约2分钟）", first_indent=False)
add_normal(doc, "好，今天这节课我们学习了生物进化的历程，知道了化石是研究生物进化最直接、最重要的证据，也学会了用比较的方法分析进化关系，还梳理了植物、无脊椎动物和脊椎动物的大致进化历程。")
add_normal(doc, "课后请大家完成两项任务：第一，整理本节课笔记；第二，自己画出植物、无脊椎动物和脊椎动物的进化简图。")
add_normal(doc, "好，这节课就上到这里。同学们辛苦了，下课！")

doc.save(str(OUT))
print(OUT)
